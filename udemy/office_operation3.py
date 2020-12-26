import docx

#document = docx.Document()
#document.save('word_sample.docx')

document = docx.Document('word_sample.docx')
document.add_heading('Document Title', 0)
p = document.add_paragraph('A PLAIN PARAGRAPH HAVING SOME\n')
p.add_run('bold\n').bold = True
p.add_run(' and more \n')
p.add_run('italic.').italic = True

document.save('word_sample.docx')